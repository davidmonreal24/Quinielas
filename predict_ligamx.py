#!/usr/bin/env python3
"""
predict_ligamx.py  —  Data Analysis Picks
==========================================
Liga MX Clausura 2026 — Predicciones + Value Bets + Documento de Metodología

Fuente de datos : data/sofascore_events.parquet  (generado por collect_sofascore.py)
Odds            : The Odds API  (https://the-odds-api.com — 500 req/mes gratuitos)
                  Registra tu clave en https://the-odds-api.com y pásala con --odds-key
Modelo          : Ratings multiplicativos ataque/defensa (Dixon-Coles inspirado)
                  + Poisson independiente + suavizado de empate (floor 15%)
Features        :
  - Forma reciente ponderada (últimos 5 partidos, más reciente = mayor peso)
  - Rendimiento local/visitante (att_h, def_h, att_a, def_a)
  - Posición en tabla acumulada (pts, gf, ga, pj)
  - H2H histórico (últimos 5 enfrentamientos directos)
  - Semáforo de confianza: ALTA / MEDIA / BAJA (basado en EV si hay momios)
EV (Value Bet)  : EV = (P_modelo × momio_decimal) − 1
                  Requiere --odds-key (The Odds API) o --odds-file (JSON manual)
Salida          :
  data/ligamx_predicciones.csv      — Predicciones tabulares (47 columnas + EV)
  data/ligamx_odds_cache.json       — Cache de momios (evita re-consultas)
  data/ligamx_metodologia.docx      — Documento Word con metodología + picks narrados

Uso             :
  python predict_ligamx.py
  python predict_ligamx.py --odds-key TU_CLAVE_ODDS_API
  python predict_ligamx.py --odds-file odds.json
"""

import argparse
import json
import math
import re
import time
from datetime import date, datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path

import numpy as np
import pandas as pd
import requests  # noqa: F401 — usado en fetch_odds_ligamx

from context_enricher import (
    dixon_coles_probs,
    smooth_draw as smooth_draw_dc,
    build_ligamx_table_context,
    apply_motivation,
    scorers_for_team,
    estimate_corners,
    match_narrative,
)

# ─── Constantes ───────────────────────────────────────────────────────────────
LIGA_MX_ID      = 11620
MIN_GAMES       = 3        # mínimo partidos para usar ratings
WINDOW          = 8        # partidos para media ponderada
DRAW_FLOOR      = 0.15     # mínimo probabilidad empate
DATA_DIR        = Path("data")
EVENTS_PATH     = DATA_DIR / "sofascore_events.parquet"
UPCOMING_PATH   = DATA_DIR / "sofascore_upcoming.parquet"
OUT_CSV         = DATA_DIR / "ligamx_predicciones.csv"
OUT_DOCX        = DATA_DIR / "ligamx_metodologia.docx"

# Umbrales de confianza (sin momios: basado en margen de probabilidad)
ALTA_GAP        = 0.18   # diferencia entre mejor y 2a probabilidad
MEDIA_GAP       = 0.10

# Umbrales EV (cuando hay momios disponibles)
EV_ALTA         = 0.10
EV_MEDIA        = 0.05

# The Odds API
ODDS_API_BASE   = "https://api.the-odds-api.com/v4"
ODDS_SPORT      = "soccer_mexico_ligamx"   # Liga MX en The Odds API
ODDS_CACHE_PATH = DATA_DIR / "ligamx_odds_cache.json"
ODDS_CACHE_TTL  = 3600   # segundos — re-fetches si el cache tiene más de 1 hora

# Bookmakers de referencia por región (The Odds API)
# 'eu' incluye: Bet365, Betway, Unibet, William Hill, etc.
ODDS_REGIONS    = "eu,us2"   # eu=europeos, us2=Pinnacle + DraftKings

# FBref — Corners y Tarjetas (via soccerdata)
FBREF_CACHE_DIR = DATA_DIR / "_fbref_cache"
FBREF_CACHE_TTL = 43200   # 12 horas en segundos
FBREF_SEASON    = "2526"  # temporada actual Liga MX

# Aliases The Odds API → nombre canónico (para fuzzy matching con Sofascore)
_ODDS_API_ALIASES = {
    "chivas guadalajara":          "chivas",
    "guadalajara":                 "chivas",
    "cd guadalajara":              "chivas",
    "club america":                "america",
    "america":                     "america",
    "cf monterrey":                "monterrey",
    "monterrey":                   "monterrey",
    "tigres uanl":                 "tigres",
    "tigres":                      "tigres",
    "cruz azul":                   "cruz azul",
    "deportivo toluca":            "toluca",
    "toluca":                      "toluca",
    "santos laguna":               "santos",
    "santos":                      "santos",
    "atlas":                       "atlas",
    "atlas fc":                    "atlas",
    "pumas unam":                  "pumas",
    "pumas":                       "pumas",
    "club leon":                   "leon",
    "leon":                        "leon",
    "necaxa":                      "necaxa",
    "club necaxa":                 "necaxa",
    "puebla":                      "puebla",
    "club puebla":                 "puebla",
    "fc juarez":                   "juarez",
    "juarez":                      "juarez",
    "atletico san luis":           "san luis",
    "atletico de san luis":        "san luis",
    "pachuca":                     "pachuca",
    "cf pachuca":                  "pachuca",
    "queretaro":                   "queretaro",
    "queretaro fc":                "queretaro",
    "mazatlan":                    "mazatlan",
    "mazatlan fc":                 "mazatlan",
    "club tijuana":                "tijuana",
    "tijuana":                     "tijuana",
}

# Aliases de equipos Liga MX para fuzzy matching
_ALIASES_LMX = {
    "club america": "america",
    "cf america": "america",
    "america fc": "america",
    "guadalajara": "chivas",
    "club deportivo guadalajara": "chivas",
    "deportivo guadalajara": "chivas",
    "tigres uanl": "tigres",
    "tigres u.a.n.l.": "tigres",
    "cf monterrey": "monterrey",
    "rayados": "monterrey",
    "club de futbol monterrey": "monterrey",
    "toluca fc": "toluca",
    "deportivo toluca": "toluca",
    "santos laguna": "santos",
    "cruz azul fc": "cruz azul",
    "deportivo cruz azul": "cruz azul",
    "pumas unam": "pumas",
    "club universidad nacional": "pumas",
    "universidad nacional": "pumas",
    "atlas fc": "atlas",
    "atlas de guadalajara": "atlas",
    "leon fc": "leon",
    "club leon": "leon",
    "necaxa fc": "necaxa",
    "club necaxa": "necaxa",
    "puebla fc": "puebla",
    "club puebla": "puebla",
    "tijuana fc": "tijuana",
    "xolos": "tijuana",
    "club tijuana": "tijuana",
    "fc juarez": "juarez",
    "bravos de juarez": "juarez",
    "atletico de san luis": "san luis",
    "atletico san luis": "san luis",
    "pachuca cf": "pachuca",
    "cf pachuca": "pachuca",
    "queretaro fc": "queretaro",
    "mazatlan fc": "mazatlan",
}


# ─── Normalización y fuzzy matching ────────────────────────────────────────────
def _norm(name: str) -> str:
    name = str(name).lower().strip()
    name = re.sub(r"\b(fc|cf|sc|ac|as|ss|ssc|afc|bv|sv|fk|sk|ud|rcd|cd|club)\b", "", name)
    name = re.sub(r"\s+", " ", name).strip()
    return _ALIASES_LMX.get(name, name)


def _best_match(target: str, candidates: list, threshold: float = 0.68) -> str | None:
    t = _norm(target)
    best, score = None, threshold
    for c in candidates:
        r = SequenceMatcher(None, t, _norm(str(c))).ratio()
        if r > score:
            score, best = r, c
    return best


def _norm_odds(name: str) -> str:
    """Normaliza nombre de equipo desde The Odds API para fuzzy matching."""
    name = str(name).lower().strip()
    name = re.sub(r"\b(fc|cf|sc|ac|cd|club|deportivo|de)\b", "", name)
    name = re.sub(r"\s+", " ", name).strip()
    return _ODDS_API_ALIASES.get(name, name)


# ─── The Odds API — Integración de momios ─────────────────────────────────────
def _load_odds_cache() -> dict:
    """Carga cache de momios si existe y es reciente (< ODDS_CACHE_TTL segundos)."""
    if not ODDS_CACHE_PATH.exists():
        return {}
    try:
        raw = json.loads(ODDS_CACHE_PATH.read_bytes().decode("utf-8"))
        ts  = raw.get("_ts", 0)
        age = datetime.now(tz=timezone.utc).timestamp() - ts
        if age < ODDS_CACHE_TTL:
            return raw.get("odds", {})
        return {}   # cache expirado
    except Exception:
        return {}


def _save_odds_cache(odds: dict) -> None:
    payload = {
        "_ts":  datetime.now(tz=timezone.utc).timestamp(),
        "odds": odds,
    }
    ODDS_CACHE_PATH.write_bytes(
        json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    )


def fetch_odds_ligamx(api_key: str) -> dict:
    """
    Obtiene momios 1X2 de Liga MX desde The Odds API.

    Clave gratuita (500 req/mes) en: https://the-odds-api.com

    Retorna: { "HomeTeam_vs_AwayTeam": {"local": float, "empate": float, "visitante": float,
                                         "bookmakers": int, "fuente": str} }
    Los momios son promedio decimal sobre todos los bookmakers disponibles.
    """
    # Intentar cache primero
    cached = _load_odds_cache()
    if cached:
        print(f"  Odds: {len(cached)} partidos desde cache "
              f"(< {ODDS_CACHE_TTL//60} min)")
        return cached

    print(f"  Consultando The Odds API ({ODDS_SPORT})...")
    try:
        r = requests.get(
            f"{ODDS_API_BASE}/sports/{ODDS_SPORT}/odds/",
            params={
                "apiKey":      api_key,
                "regions":     ODDS_REGIONS,
                "markets":     "h2h",          # 1X2
                "oddsFormat":  "decimal",
                "dateFormat":  "iso",
            },
            timeout=15,
        )
    except Exception as e:
        print(f"  AVISO odds: error de red — {e}")
        return {}

    remaining  = r.headers.get("x-requests-remaining", "?")
    used       = r.headers.get("x-requests-used", "?")

    if r.status_code == 401:
        print("  AVISO odds: clave inválida. Verifica en https://the-odds-api.com")
        return {}
    if r.status_code == 422:
        print(f"  AVISO odds: deporte '{ODDS_SPORT}' no disponible con tu plan.")
        return {}
    if r.status_code != 200:
        print(f"  AVISO odds: HTTP {r.status_code} — {r.text[:200]}")
        return {}

    print(f"  Odds API: {used} usadas / {remaining} restantes este mes")
    events = r.json()
    print(f"  Partidos con odds disponibles: {len(events)}")

    odds_map = {}
    for event in events:
        home_api  = event.get("home_team", "")
        away_api  = event.get("away_team", "")

        # Acumular odds por resultado (promedio entre bookmakers)
        acc = {"1": [], "X": [], "2": []}
        bk_names = []
        for bookmaker in event.get("bookmakers", []):
            bk_names.append(bookmaker.get("title", ""))
            for market in bookmaker.get("markets", []):
                if market.get("key") != "h2h":
                    continue
                for outcome in market.get("outcomes", []):
                    nm    = outcome["name"]
                    price = float(outcome["price"])
                    if nm == home_api:
                        acc["1"].append(price)
                    elif nm == away_api:
                        acc["2"].append(price)
                    else:   # Draw / Empate
                        acc["X"].append(price)

        if not acc["1"] or not acc["2"]:
            continue   # sin datos suficientes

        avg_1 = round(sum(acc["1"]) / len(acc["1"]), 3)
        avg_x = round(sum(acc["X"]) / len(acc["X"]), 3) if acc["X"] else None
        avg_2 = round(sum(acc["2"]) / len(acc["2"]), 3)

        # Probabilidades implícitas (ajustadas por vig)
        p_sum = 1/avg_1 + (1/avg_x if avg_x else 0) + 1/avg_2
        p_imp_1 = round((1/avg_1) / p_sum, 4) if p_sum else None
        p_imp_x = round((1/avg_x) / p_sum, 4) if avg_x and p_sum else None
        p_imp_2 = round((1/avg_2) / p_sum, 4) if p_sum else None

        key = f"{home_api}_vs_{away_api}"
        odds_map[key] = {
            "local":      avg_1,
            "empate":     avg_x,
            "visitante":  avg_2,
            "p_imp_local":   p_imp_1,
            "p_imp_empate":  p_imp_x,
            "p_imp_visit":   p_imp_2,
            "bookmakers": len(acc["1"]),
            "casas":      ", ".join(dict.fromkeys(bk_names))[:80],
            "home_api":   home_api,
            "away_api":   away_api,
        }

    _save_odds_cache(odds_map)
    return odds_map


# ─── FBref Liga MX: Corners y Tarjetas ────────────────────────────────────────

def _load_apifootball_liga_mx_cards(seasons: tuple = (2025, 2024)) -> dict:
    """
    Obtiene promedios de tarjetas amarillas desde API-Football.
    Endpoint: GET /teams/statistics?league=262&season={s}&team={id}
    Retorna: {team_norm: {"avg_yellows": float, "n_games": int}}
    """
    from context_enricher import APIFOOTBALL_BASE, APIFOOTBALL_HDR

    for season_id in seasons:
        try:
            r = requests.get(
                f"{APIFOOTBALL_BASE}/teams",
                headers=APIFOOTBALL_HDR,
                params={"league": 262, "season": season_id},
                timeout=10,
            )
            if r.status_code != 200:
                continue
            teams = r.json().get("response", [])
            if not teams:
                continue

            print(f"  API-Football Liga MX {season_id}: {len(teams)} equipos")
            result: dict = {}
            for td in teams[:20]:
                tid   = td["team"]["id"]
                tname = td["team"]["name"]
                time.sleep(0.4)
                sr = requests.get(
                    f"{APIFOOTBALL_BASE}/teams/statistics",
                    headers=APIFOOTBALL_HDR,
                    params={"league": 262, "season": season_id, "team": tid},
                    timeout=10,
                )
                if sr.status_code != 200:
                    continue
                ts = sr.json().get("response", {})
                yellow_data = ts.get("cards", {}).get("yellow", {})
                total_yellows = sum((v.get("total") or 0) for v in yellow_data.values())
                games_total = ts.get("fixtures", {}).get("played", {}).get("total", 0) or 1
                tn = _norm(tname)
                result[tn] = {
                    "avg_yellows": round(total_yellows / games_total, 1),
                    "n_games":     games_total,
                }
            if result:
                print(f"  API-Football: {len(result)} equipos con tarjetas (season {season_id})")
                return result
        except Exception as ex:
            print(f"  API-Football ({season_id}): {ex}")
    return {}


def load_fbref_ligamx_stats(season: str = FBREF_SEASON) -> dict:
    """
    Descarga estadísticas históricas de corners y tarjetas para Liga MX.

    Estrategia en cascada:
      1. FBref via soccerdata — passing_types (CK) y misc (CrdY).
         Nota: soccerdata solo soporta Big-5; falla para Liga MX y se captura.
      2. Fallback: API-Football /teams/statistics — tarjetas amarillas por equipo.
         (Corners: se estiman con att_h-calibration en generate_predictions)

    Retorna: {
      "nombre_normalizado": {
        "avg_corners":   float | None,
        "avg_corners_h": float | None,
        "avg_corners_a": float | None,
        "avg_yellows":   float | None,
        "avg_yellows_h": float | None,
        "avg_yellows_a": float | None,
        "n_games":       int,
      }
    }
    """
    FBREF_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_file = FBREF_CACHE_DIR / f"ligamx_{season}_stats.json"

    if cache_file.exists():
        age = datetime.now(tz=timezone.utc).timestamp() - cache_file.stat().st_mtime
        if age < FBREF_CACHE_TTL:
            try:
                data = json.loads(cache_file.read_bytes().decode("utf-8"))
                if data is not None:
                    print(f"  Caché stats: {len(data)} equipos")
                    return data
            except Exception:
                pass

    stats: dict = {}

    # ── Intento 1: FBref via soccerdata ─────────────────────────────────────────
    try:
        import soccerdata as sd
        raw_dir = str(DATA_DIR / "_fbref_raw")

        def _col(df: pd.DataFrame, *names: str) -> str | None:
            col_map = {c.lower(): c for c in df.columns}
            for n in names:
                if n.lower() in col_map:
                    return col_map[n.lower()]
            return None

        def _team_avgs_fbref(df: pd.DataFrame, stat: str, prefix: str) -> None:
            tc = _col(df, "team", "squad", "club")
            vc = _col(df, "home_away", "venue", "ha", "h_a")
            sc = _col(df, stat)
            if not tc or not sc:
                return
            for team, grp in df.groupby(tc):
                tn = _norm(str(team))
                if tn not in stats:
                    stats[tn] = {}
                vals = pd.to_numeric(grp[sc], errors="coerce").dropna()
                if len(vals) == 0:
                    continue
                stats[tn][f"avg_{prefix}"] = round(float(vals.mean()), 1)
                stats[tn]["n_games"] = max(stats[tn].get("n_games", 0), len(vals))
                if vc:
                    h = grp[grp[vc].astype(str).str.upper().str.startswith("H")]
                    a = grp[grp[vc].astype(str).str.upper().str.startswith("A")]
                    vh = pd.to_numeric(h[sc], errors="coerce").dropna()
                    va = pd.to_numeric(a[sc], errors="coerce").dropna()
                    stats[tn][f"avg_{prefix}_h"] = round(float(vh.mean()), 1) if len(vh) > 0 else None
                    stats[tn][f"avg_{prefix}_a"] = round(float(va.mean()), 1) if len(va) > 0 else None

        fbref = sd.FBref(leagues=["MEX-Liga MX"], seasons=[season], data_dir=raw_dir)
        try:
            pt = fbref.read_team_match_stats(stat_type="passing_types")
            if pt is not None and not pt.empty:
                idx = pt.index
                df_pt = pt.reset_index() if (idx.name or (hasattr(idx, "names") and idx.names[0])) else pt.copy()
                _team_avgs_fbref(df_pt, "CK", "corners")
        except Exception:
            pass
        try:
            mc = fbref.read_team_match_stats(stat_type="misc")
            if mc is not None and not mc.empty:
                idx = mc.index
                df_mc = mc.reset_index() if (idx.name or (hasattr(idx, "names") and idx.names[0])) else mc.copy()
                _team_avgs_fbref(df_mc, "CrdY", "yellows")
        except Exception:
            pass
        if stats:
            print(f"  FBref: {len(stats)} equipos Liga MX OK")
    except Exception:
        pass  # Liga MX no soportada por soccerdata-FBref; caer en fallback

    # ── Intento 2: API-Football (tarjetas amarillas) ──────────────────────────────
    if not any("avg_yellows" in v for v in stats.values()):
        print("  FBref sin datos Liga MX -> fallback API-Football (tarjetas)...")
        af = _load_apifootball_liga_mx_cards()
        for tn, v in af.items():
            if tn not in stats:
                stats[tn] = {}
            stats[tn].update(v)

    cache_file.write_bytes(json.dumps(stats, ensure_ascii=False, indent=2).encode("utf-8"))
    if stats:
        has_yel = sum(1 for v in stats.values() if v.get("avg_yellows") is not None)
        has_ck  = sum(1 for v in stats.values() if v.get("avg_corners") is not None)
        print(f"  Stats: {len(stats)} equipos | corners={has_ck} | amarillas={has_yel}")
    else:
        print("  Sin datos — corners estimados por λ-att, amarillas N/D")
    return stats


def _get_fbref(team: str, fbref_stats: dict) -> dict:
    """Busca stats por exacto, contenido parcial, y fuzzy con umbral reducido."""
    if not fbref_stats:
        return {}
    tn = _norm(team)
    # 1. Exacto
    if tn in fbref_stats:
        return fbref_stats[tn]
    # 2. Contenido: tn en clave o clave en tn (ej. "guadalajara" en "guadalajara chivas")
    for key in fbref_stats:
        if (tn and key and tn in key) or (key and tn and key in tn):
            return fbref_stats[key]
    # 3. Fuzzy con umbral reducido a 0.55
    matched = _best_match(team, list(fbref_stats.keys()), threshold=0.55)
    return fbref_stats.get(matched, {}) if matched else {}


def match_odds_to_fixture(
    home: str,
    away: str,
    odds_map: dict,
) -> dict | None:
    """
    Busca los momios del partido en odds_map usando fuzzy matching.
    Prueba primero clave exacta, luego normalización, luego SequenceMatcher.
    """
    # Clave exacta por nombre normalizado
    for v in odds_map.values():
        h_api = v.get("home_api", "")
        a_api = v.get("away_api", "")
        if (SequenceMatcher(None, _norm_odds(home), _norm_odds(h_api)).ratio() > 0.75
                and SequenceMatcher(None, _norm_odds(away), _norm_odds(a_api)).ratio() > 0.75):
            return v

    # Fuzzy sobre todos los partidos
    t_home = _norm_odds(home)
    t_away = _norm_odds(away)
    best_score, best_val = 0.0, None
    for v in odds_map.values():
        h_score = SequenceMatcher(None, t_home, _norm_odds(v.get("home_api", ""))).ratio()
        a_score = SequenceMatcher(None, t_away, _norm_odds(v.get("away_api", ""))).ratio()
        combined = (h_score + a_score) / 2
        if combined > best_score:
            best_score = combined
            best_val = v

    return best_val if best_score > 0.65 else None


# ─── Utilidades estadísticas ───────────────────────────────────────────────────
def _weighted_mean(series: pd.Series) -> float:
    n = len(series)
    if n == 0:
        return 0.0
    weights = np.arange(1, n + 1, dtype=float)
    return float(np.average(series.values, weights=weights))


def _clip(v, lo: float = 0.30, hi: float = 2.20) -> float:
    """Clip individual rating factor. Reducido a 2.2 para evitar lambdas extremas."""
    if v is None:
        return 1.0
    return max(lo, min(hi, float(v)))


LAMBDA_MAX = 3.2   # Ningún equipo de Liga MX debería esperar más de 3.2 goles

def _clip_lambda(lam: float) -> float:
    return max(0.20, min(LAMBDA_MAX, lam))


def _form_multiplier(forma: str) -> float:
    """Factor 0.92–1.08 basado en últimos 3 resultados recientes."""
    if not forma:
        return 1.0
    recent3 = forma[-3:]
    pts = sum({"W": 3, "D": 1, "L": 0}.get(c, 0) for c in recent3)
    max_pts = len(recent3) * 3
    # Escala lineal: 0 pts→0.92, 3pts→1.0, 6pts→1.04, 9pts→1.08
    return round(0.92 + (pts / max(max_pts, 1)) * 0.16, 3)


def _ppmf(k: int, lam: float) -> float:
    if lam <= 0 or math.isnan(lam):
        return 1.0 if k == 0 else 0.0
    return float(math.exp(-lam) * (lam ** k) / math.factorial(k))


def poisson_probs(lh: float, la: float, max_g: int = 9) -> tuple:
    pl = pd_ = pv = 0.0
    for h in range(max_g + 1):
        for a in range(max_g + 1):
            p = _ppmf(h, lh) * _ppmf(a, la)
            if h > a:     pl  += p
            elif h == a:  pd_ += p
            else:         pv  += p
    total = pl + pd_ + pv
    return (pl / total, pd_ / total, pv / total) if total else (1/3, 1/3, 1/3)


def smooth_draw(pl: float, pd_: float, pv: float, floor: float = DRAW_FLOOR) -> tuple:
    if pd_ >= floor:
        return pl, pd_, pv
    deficit = floor - pd_
    lv = pl + pv
    if lv <= 0:
        return pl, floor, pv
    return pl - deficit * (pl / lv), floor, pv - deficit * (pv / lv)


def form_pts(forma: str) -> int:
    return sum({"W": 3, "D": 1, "L": 0}.get(c, 0) for c in forma)


# ─── Ratings multiplicativos ────────────────────────────────────────────────────
# Shrinkage Bayesiano: con pocos juegos, el rating se acerca a 1.0 (promedio liga)
# Formula: rating_adj = (n * rating + k) / (n + k)  donde k = prior strength
SHRINK_K = 4   # 4 partidos = mitad del camino entre rating puro y promedio

def _shrink(rating: float, n: int, k: float = SHRINK_K) -> float:
    """Regresa rating hacia 1.0 cuando hay pocos partidos (shrinkage Bayesiano)."""
    return (n * rating + k * 1.0) / (n + k)


def _exp_weighted_mean(series: pd.Series, alpha: float = 0.7) -> float:
    """Media ponderada exponencial: más reciente = mayor peso (EWM)."""
    n = len(series)
    if n == 0:
        return 0.0
    # Pesos: alpha^(n-1), alpha^(n-2), ..., alpha^0 (más reciente = alpha^0 = 1)
    weights = np.array([alpha ** (n - 1 - i) for i in range(n)], dtype=float)
    return float(np.average(series.values, weights=weights))


def compute_ratings(df: pd.DataFrame) -> dict:
    """
    Calcula ratings multiplicativos att/def para cada equipo.
    Mejoras v2: shrinkage Bayesiano, ponderación exponencial, ventana deslizante corta.
    Retorna: {"teams": {...}, "mu_h": float, "mu_a": float, "home_adv": float}
    """
    fallback = {"teams": {}, "mu_h": 1.40, "mu_a": 1.18, "home_adv": 1.19}
    if df is None or df.empty:
        return fallback

    df = df.dropna(subset=["home_goals", "away_goals"]).copy()
    if len(df) < 5:
        return fallback

    mu_h = float(df["home_goals"].mean())
    mu_a = float(df["away_goals"].mean())
    if mu_h <= 0 or mu_a <= 0:
        return fallback

    home_adv = mu_h / mu_a
    teams_all = pd.concat([df["home_team"], df["away_team"]]).unique()
    ratings = {}

    for team in teams_all:
        hg = df[df["home_team"] == team].sort_values("date")
        ag = df[df["away_team"] == team].sort_values("date")
        nh, na = len(hg), len(ag)

        # Ponderación exponencial con ventana deslizante
        hg_w = hg.tail(WINDOW)
        ag_w = ag.tail(WINDOW)
        nh_w = len(hg_w)
        na_w = len(ag_w)

        # Rating bruto con EWM
        if nh_w >= MIN_GAMES:
            raw_att_h = _exp_weighted_mean(hg_w["home_goals"]) / mu_h
            raw_def_h = _exp_weighted_mean(hg_w["away_goals"]) / mu_a
            att_h = _shrink(raw_att_h, nh_w)
            def_h = _shrink(raw_def_h, nh_w)
        else:
            att_h = def_h = None

        if na_w >= MIN_GAMES:
            raw_att_a = _exp_weighted_mean(ag_w["away_goals"]) / mu_a
            raw_def_a = _exp_weighted_mean(ag_w["home_goals"]) / mu_h
            att_a = _shrink(raw_att_a, na_w)
            def_a = _shrink(raw_def_a, na_w)
        else:
            att_a = def_a = None

        # Forma reciente (últimos 5 partidos, orden cronológico)
        recent = (
            pd.concat([
                hg[["date", "home_goals", "away_goals"]]
                  .rename(columns={"home_goals": "gf", "away_goals": "ga"}),
                ag[["date", "away_goals", "home_goals"]]
                  .rename(columns={"away_goals": "gf", "home_goals": "ga"}),
            ])
            .sort_values("date")
            .tail(5)
        )
        forma = "".join(
            "W" if float(r["gf"]) > float(r["ga"])
            else ("D" if float(r["gf"]) == float(r["ga"]) else "L")
            for _, r in recent.iterrows()
        )

        ratings[team] = {
            "att_h": att_h, "def_h": def_h,
            "att_a": att_a, "def_a": def_a,
            "nh": nh,       "na": na,
            "forma": forma,
        }

    return {"teams": ratings, "mu_h": mu_h, "mu_a": mu_a, "home_adv": home_adv}


# ─── Tabla de posiciones ─────────────────────────────────────────────────────────
def compute_table(df: pd.DataFrame) -> pd.DataFrame:
    """
    Calcula la tabla de posiciones desde los resultados históricos.
    Usa la temporada más reciente disponible.
    """
    df = df.dropna(subset=["home_goals", "away_goals"]).copy()
    if df.empty:
        return pd.DataFrame()

    # Filtrar por temporada más reciente si hay datos de season_year
    if "season_year" in df.columns:
        seasons = df["season_year"].dropna().unique()
        if len(seasons) > 0:
            latest = sorted(str(s) for s in seasons if s)[-1]
            df_season = df[df["season_year"].astype(str) == latest]
            if len(df_season) >= 5:
                df = df_season

    records = []
    for _, row in df.iterrows():
        hg, ag = int(row["home_goals"]), int(row["away_goals"])
        h_pts = 3 if hg > ag else (1 if hg == ag else 0)
        a_pts = 3 if ag > hg else (1 if ag == hg else 0)
        records.append({
            "team": row["home_team"], "pts": h_pts, "pj": 1,
            "wins": 1 if hg > ag else 0, "draws": 1 if hg == ag else 0,
            "losses": 1 if hg < ag else 0, "gf": hg, "ga": ag,
        })
        records.append({
            "team": row["away_team"], "pts": a_pts, "pj": 1,
            "wins": 1 if ag > hg else 0, "draws": 1 if ag == hg else 0,
            "losses": 1 if ag < hg else 0, "gf": ag, "ga": hg,
        })

    if not records:
        return pd.DataFrame()

    table = pd.DataFrame(records).groupby("team").sum().reset_index()
    table["gd"] = table["gf"] - table["ga"]
    table = (
        table
        .sort_values(["pts", "gd", "gf"], ascending=False)
        .reset_index(drop=True)
    )
    table["position"] = table.index + 1
    return table.set_index("team")


def _table_info(table: pd.DataFrame, team: str) -> dict:
    """Busca la posición en tabla del equipo (con fuzzy matching)."""
    empty = {"pos": None, "pts_tabla": None, "pj": None}
    if table is None or table.empty:
        return empty

    # Búsqueda directa
    if team in table.index:
        r = table.loc[team]
        return {"pos": int(r["position"]), "pts_tabla": int(r["pts"]), "pj": int(r["pj"])}

    # Fuzzy
    matched = _best_match(team, list(table.index))
    if matched:
        r = table.loc[matched]
        return {"pos": int(r["position"]), "pts_tabla": int(r["pts"]), "pj": int(r["pj"])}
    return empty


# ─── H2H ─────────────────────────────────────────────────────────────────────────
def compute_h2h(df: pd.DataFrame, home: str, away: str, n: int = 5) -> dict:
    """H2H entre dos equipos (últimos n enfrentamientos, ambas direcciones)."""
    mask = (
        ((df["home_team"] == home) & (df["away_team"] == away)) |
        ((df["home_team"] == away) & (df["away_team"] == home))
    )
    h2h = (
        df[mask]
        .dropna(subset=["home_goals", "away_goals"])
        .sort_values("date")
        .tail(n)
    )

    if h2h.empty:
        return {"h2h_n": 0, "h2h_w_h": 0, "h2h_d": 0, "h2h_w_a": 0,
                "h2h_gf_h": 0, "h2h_gf_a": 0}

    w_h = d = w_a = gf_h = gf_a = 0
    for _, row in h2h.iterrows():
        hg, ag = int(row["home_goals"]), int(row["away_goals"])
        if row["home_team"] == home:
            gf_h += hg; gf_a += ag
            if hg > ag:   w_h += 1
            elif hg == ag: d   += 1
            else:          w_a += 1
        else:
            gf_h += ag; gf_a += hg
            if ag > hg:   w_h += 1
            elif ag == hg: d   += 1
            else:          w_a += 1

    return {"h2h_n": len(h2h), "h2h_w_h": w_h, "h2h_d": d, "h2h_w_a": w_a,
            "h2h_gf_h": gf_h, "h2h_gf_a": gf_a}


# ─── Semáforo de confianza ────────────────────────────────────────────────────────
def semaforo(
    p_loc: float, p_emp: float, p_vis: float,
    ev_max: float | None = None,
) -> str:
    """
    ALTA / MEDIA / BAJA:
      - Si hay EV: ALTA si ev_max > 0.10, MEDIA si 0.05-0.10, BAJA si < 0.05
      - Sin EV: basado en margen entre prob. máxima y segunda
    """
    if ev_max is not None:
        if ev_max >= EV_ALTA:   return "ALTA"
        if ev_max >= EV_MEDIA:  return "MEDIA"
        return "BAJA"

    probs = sorted([p_loc, p_emp, p_vis], reverse=True)
    gap = probs[0] - probs[1]
    if gap >= ALTA_GAP:   return "ALTA"
    if gap >= MEDIA_GAP:  return "MEDIA"
    return "BAJA"


# ─── EV (Value Bet) ────────────────────────────────────────────────────────────────
def calc_ev(p_modelo: float, momio_decimal: float) -> float:
    """EV = (P_modelo × momio) − 1. Positivo = value bet."""
    return round(p_modelo * momio_decimal - 1, 4)


# ─── Predicciones ─────────────────────────────────────────────────────────────────
def _get_team_ratings(team: str, ratings: dict) -> dict:
    teams = ratings.get("teams", {})
    if team in teams:
        return teams[team]
    matched = _best_match(team, list(teams.keys()))
    return teams.get(matched, {}) if matched else {}


def generate_predictions(
    upcoming: list,
    hist_df: pd.DataFrame,
    ratings: dict,
    table: pd.DataFrame,
    odds: dict | None = None,
    fbref_stats: dict | None = None,
) -> pd.DataFrame:
    """Genera predicciones Liga MX con todos los features requeridos."""
    mu_h = ratings["mu_h"]
    mu_a = ratings["mu_a"]
    rows = []

    # Construir contexto de motivación desde tabla
    # compute_table usa el nombre del equipo como índice y columna "position"
    table_rows = []
    for team_name, tr in table.iterrows():
        table_rows.append({
            "team": team_name,
            "pos":  int(tr.get("position", tr.get("pos", 99))),
            "pts":  int(tr.get("pts", 0)),
        })
    table_rows_sorted = sorted(table_rows, key=lambda x: x["pos"])
    motiv_ctx = build_ligamx_table_context(table_rows_sorted)
    print(f"  Contexto de motivación: {len(motiv_ctx)} equipos")

    # Pre-cargar goleadores (una sola llamada API)
    print("  Cargando goleadores (API-Football)...")
    _ = scorers_for_team("placeholder", "ligamx")  # warm cache

    for rec in sorted(upcoming, key=lambda x: x.get("date", "")):
        home  = rec["home_team"]
        away  = rec["away_team"]
        m_date = str(rec.get("date", ""))[:10]
        fase  = rec.get("round_name", "")

        h_r = _get_team_ratings(home, ratings)
        a_r = _get_team_ratings(away, ratings)

        att_h = h_r.get("att_h")
        def_h = h_r.get("def_h")
        att_a = a_r.get("att_a")
        def_a = a_r.get("def_a")

        lh_base = _clip_lambda(_clip(att_h) * _clip(def_a) * mu_h)
        la_base = _clip_lambda(_clip(att_a) * _clip(def_h) * mu_a)

        # ── Forma reciente como multiplicador (±8% basado en últimos 3) ──
        forma_h = h_r.get("forma", "")
        forma_a = a_r.get("forma", "")
        lh_base = _clip_lambda(lh_base * _form_multiplier(forma_h))
        la_base = _clip_lambda(la_base * _form_multiplier(forma_a))

        # ── Motivación ──
        ctx_h = motiv_ctx.get(home) or motiv_ctx.get(
            _best_match(home, list(motiv_ctx.keys())) or home)
        ctx_a = motiv_ctx.get(away) or motiv_ctx.get(
            _best_match(away, list(motiv_ctx.keys())) or away)
        f_h = ctx_h["factor"] if ctx_h else 1.0
        f_a = ctx_a["factor"] if ctx_a else 1.0
        lh, la = apply_motivation(lh_base, la_base, f_h, f_a)
        lh, la = _clip_lambda(lh), _clip_lambda(la)

        # ── Dixon-Coles (mejor estimación de empate) ──
        p_loc_r, p_emp_r, p_vis_r = dixon_coles_probs(lh, la)
        p_loc, p_emp, p_vis       = smooth_draw_dc(p_loc_r, p_emp_r, p_vis_r)

        probs  = [p_loc, p_emp, p_vis]
        labels = ["Local", "Empate", "Visitante"]
        best_idx = probs.index(max(probs))
        # Reportar "Empate" cuando P(X)>25% y la diferencia con el mejor es <10pp
        sorted_p = sorted(probs, reverse=True)
        if best_idx != 1 and p_emp >= 0.25 and (sorted_p[0] - p_emp) < 0.10:
            pred = "Empate"
        else:
            pred = labels[best_idx]

        # Odds y EV — fuzzy matching contra The Odds API
        ev_loc = ev_emp = ev_vis = ev_max = None
        momio_loc = momio_emp = momio_vis = None
        p_imp_loc = p_imp_emp = p_imp_vis = None
        n_bookmakers = None
        casas_ref = None

        if odds:
            o = match_odds_to_fixture(home, away, odds)
            if o:
                momio_loc    = o.get("local")
                momio_emp    = o.get("empate")
                momio_vis    = o.get("visitante")
                p_imp_loc    = o.get("p_imp_local")
                p_imp_emp    = o.get("p_imp_empate")
                p_imp_vis    = o.get("p_imp_visit")
                n_bookmakers = o.get("bookmakers")
                casas_ref    = o.get("casas", "")
                if momio_loc: ev_loc = calc_ev(p_loc, momio_loc)
                if momio_emp: ev_emp = calc_ev(p_emp, momio_emp)
                if momio_vis: ev_vis = calc_ev(p_vis, momio_vis)
                valid_evs = [v for v in [ev_loc, ev_emp, ev_vis] if v is not None]
                ev_max = max(valid_evs) if valid_evs else None

        confianza = semaforo(p_loc, p_emp, p_vis, ev_max)

        # Contexto: tabla y H2H
        tbl_h = _table_info(table, home)
        tbl_a = _table_info(table, away)
        h2h   = compute_h2h(hist_df, home, away)

        # ── Contexto enriquecido ──
        corners = estimate_corners(lh, la)

        # ── Datos FBref: corners reales y amarillas históricas ──
        h_fbr = _get_fbref(home, fbref_stats or {})
        a_fbr = _get_fbref(away, fbref_stats or {})

        # Corners: FBref si disponible; si no, base fija Liga MX calibrada por att
        # NOTA: no usar corners_h_est (ya contiene att_h vía λ → doble efecto)
        MX_BASE_CK_H = 5.0   # Liga MX: ~5 corners/local/partido
        MX_BASE_CK_A = 4.5   # Liga MX: ~4.5 corners/visitante/partido
        ck_h = h_fbr.get("avg_corners_h")
        ck_a = a_fbr.get("avg_corners_a")
        if ck_h is not None and ck_a is not None:
            corners_h_pred  = round(float(ck_h), 1)
            corners_a_pred  = round(float(ck_a), 1)
            corners_total_p = round(corners_h_pred + corners_a_pred, 1)
            corners_fuente  = "FBref"
        elif att_h is not None and att_a is not None:
            # Factor att suavizado (^0.4): att=1.0→x1.0, att=2.0→x1.32, att=0.5→x0.76
            att_h_f = max(0.65, min(1.40, float(att_h) ** 0.4))
            att_a_f = max(0.65, min(1.40, float(att_a) ** 0.4))
            corners_h_pred  = round(max(2.0, MX_BASE_CK_H * att_h_f), 1)
            corners_a_pred  = round(max(2.0, MX_BASE_CK_A * att_a_f), 1)
            corners_total_p = round(corners_h_pred + corners_a_pred, 1)
            corners_fuente  = "att-calibrado"
        else:
            corners_h_pred  = corners["corners_h_est"]
            corners_a_pred  = corners["corners_a_est"]
            corners_total_p = corners["total_est"]
            corners_fuente  = "λ"

        # Tarjetas amarillas desde FBref (home vs away diferenciado)
        amar_h = h_fbr.get("avg_yellows_h") or h_fbr.get("avg_yellows")
        amar_a = a_fbr.get("avg_yellows_a") or a_fbr.get("avg_yellows")

        scorers_h = scorers_for_team(home, "ligamx")
        scorers_a = scorers_for_team(away, "ligamx")
        narr = match_narrative(
            home, away, ctx_h, ctx_a, p_loc, p_emp, p_vis, lh, la,
            h2h.get("h2h_w_h", 0), h2h.get("h2h_d", 0), h2h.get("h2h_w_a", 0),
        )

        rows.append({
            "fecha":            m_date,
            "fase":             fase,
            "local":            home,
            "visitante":        away,
            "prediccion":       pred,
            "p_local_pct":      round(p_loc * 100, 1),
            "p_empate_pct":     round(p_emp * 100, 1),
            "p_visit_pct":      round(p_vis * 100, 1),
            "lambda_h":         round(lh, 3),
            "lambda_a":         round(la, 3),
            "ventaja":          round(lh - la, 3),
            "att_h":            round(att_h, 3) if att_h else None,
            "def_h":            round(def_h, 3) if def_h else None,
            "att_a":            round(att_a, 3) if att_a else None,
            "def_a":            round(def_a, 3) if def_a else None,
            "n_h":              h_r.get("nh", 0),
            "n_a":              a_r.get("na", 0),
            "forma_h":          forma_h or "N/D",
            "forma_a":          forma_a or "N/D",
            "pts_forma_h":      form_pts(forma_h),
            "pts_forma_a":      form_pts(forma_a),
            "pos_tabla_h":      tbl_h["pos"],
            "pts_tabla_h":      tbl_h["pts_tabla"],
            "pj_h":             tbl_h["pj"],
            "pos_tabla_v":      tbl_a["pos"],
            "pts_tabla_v":      tbl_a["pts_tabla"],
            "pj_v":             tbl_a["pj"],
            "h2h_n":            h2h["h2h_n"],
            "h2h_w_h":          h2h["h2h_w_h"],
            "h2h_d":            h2h["h2h_d"],
            "h2h_w_a":          h2h["h2h_w_a"],
            "h2h_gf_h":         h2h["h2h_gf_h"],
            "h2h_gf_a":         h2h["h2h_gf_a"],
            "mu_h":             round(mu_h, 3),
            "mu_a":             round(mu_a, 3),
            "home_adv":         round(ratings["home_adv"], 3),
            "p_local_raw":      round(p_loc_r * 100, 1),
            "p_empate_raw":     round(p_emp_r * 100, 1),
            "p_visit_raw":      round(p_vis_r * 100, 1),
            "nivel_confianza":  confianza,
            "momio_ref_local":  momio_loc,
            "momio_ref_empate":  momio_emp,
            "momio_ref_visit":   momio_vis,
            "p_imp_local":       p_imp_loc,
            "p_imp_empate":      p_imp_emp,
            "p_imp_visit":       p_imp_vis,
            "n_bookmakers":      n_bookmakers,
            "casas_referencia":  casas_ref,
            "ev_local":          ev_loc,
            "ev_empate":         ev_emp,
            "ev_visitante":      ev_vis,
            "es_value_bet":      True if ev_max is not None and ev_max > 0 else None,
            # ── Nuevos campos de contexto ──
            "motivacion_local":    ctx_h["label"] if ctx_h else "N/D",
            "motivacion_icon_h":   ctx_h["icon"]  if ctx_h else "",
            "factor_motiv_h":      round(f_h, 3),
            "motivacion_visita":   ctx_a["label"] if ctx_a else "N/D",
            "motivacion_icon_a":   ctx_a["icon"]  if ctx_a else "",
            "factor_motiv_a":      round(f_a, 3),
            "alerta_empate":       narr["draw_alert"],
            "narrativa":           " | ".join(narr["narrative"]),
            "corners_h_est":       corners["corners_h_est"],
            "corners_a_est":       corners["corners_a_est"],
            "corners_total_rango": corners["total_range"],
            # FBref: corners predichos (real si disponible, else lambda)
            "corners_h_pred":      corners_h_pred,
            "corners_a_pred":      corners_a_pred,
            "corners_total_pred":  corners_total_p,
            "corners_fuente":      corners_fuente,
            # FBref: tarjetas amarillas históricas
            "amarillas_local":     round(float(amar_h), 1) if amar_h else None,
            "amarillas_visita":    round(float(amar_a), 1) if amar_a else None,
            "amarillas_total":     round(float(amar_h or 0) + float(amar_a or 0), 1) if (amar_h and amar_a) else None,
            "goleadores_local":    ", ".join(
                f"{s['player']}({s['goals']}g)" for s in scorers_h
            ) if scorers_h else "N/D (ref. temporada anterior)",
            "goleadores_visita":   ", ".join(
                f"{s['player']}({s['goals']}g)" for s in scorers_a
            ) if scorers_a else "N/D (ref. temporada anterior)",
            "lambda_h_base":       round(lh_base, 3),
            "lambda_a_base":       round(la_base, 3),
        })

    return pd.DataFrame(rows)


# ─── Nombres de columnas en lenguaje natural ─────────────────────────────────────
COL_NAMES = {
    "fecha":             "Fecha",
    "fase":              "Fase / Jornada",
    "local":             "Equipo Local",
    "visitante":         "Equipo Visitante",
    "prediccion":        "Prediccion",
    "p_local_pct":       "Probabilidad Local % (p_local)",
    "p_empate_pct":      "Probabilidad Empate % (p_empate)",
    "p_visit_pct":       "Probabilidad Visitante % (p_visit)",
    "lambda_h":          "Goles Esperados Local (lambda_h)",
    "lambda_a":          "Goles Esperados Visitante (lambda_a)",
    "ventaja":           "Margen Goles Esperado (ventaja = lambda_h - lambda_a)",
    "att_h":             "Ratio Ataque Local (att_h, >1 = mejor que promedio)",
    "def_h":             "Ratio Defensa Local (def_h, <1 = mejor que promedio)",
    "att_a":             "Ratio Ataque Visitante (att_a)",
    "def_a":             "Ratio Defensa Visitante (def_a)",
    "n_h":               "Partidos Base Local (n_h)",
    "n_a":               "Partidos Base Visitante (n_a)",
    "forma_h":           "Forma Reciente Local W/D/L (forma_h)",
    "forma_a":           "Forma Reciente Visitante W/D/L (forma_a)",
    "pts_forma_h":       "Puntos Forma Local (pts_h, W=3 D=1 L=0)",
    "pts_forma_a":       "Puntos Forma Visitante (pts_a, W=3 D=1 L=0)",
    "pos_tabla_h":       "Posicion Tabla Local (pos_h)",
    "pts_tabla_h":       "Puntos en Tabla Local (pts_tabla_h)",
    "pj_h":              "Partidos Jugados Local (pj_h)",
    "pos_tabla_v":       "Posicion Tabla Visitante (pos_v)",
    "pts_tabla_v":       "Puntos en Tabla Visitante (pts_tabla_v)",
    "pj_v":              "Partidos Jugados Visitante (pj_v)",
    "h2h_n":             "H2H Partidos Analizados (h2h_n)",
    "h2h_w_h":           "H2H Victorias Local (h2h_w_h)",
    "h2h_d":             "H2H Empates (h2h_d)",
    "h2h_w_a":           "H2H Victorias Visitante (h2h_w_a)",
    "h2h_gf_h":          "H2H Goles Favor Local (h2h_gf_h)",
    "h2h_gf_a":          "H2H Goles Favor Visitante (h2h_gf_a)",
    "mu_h":              "Media Goles Local Liga (mu_h)",
    "mu_a":              "Media Goles Visitante Liga (mu_a)",
    "home_adv":          "Factor Ventaja Local (home_adv)",
    "p_local_raw":       "Prob Local % sin suavizar (p_local_raw)",
    "p_empate_raw":      "Prob Empate % sin suavizar (p_empate_raw)",
    "p_visit_raw":       "Prob Visitante % sin suavizar (p_visit_raw)",
    "nivel_confianza":   "Nivel de Confianza (ALTA / MEDIA / BAJA)",
    "momio_ref_local":   "Momio Referencia Local (momio_ref_local)",
    "momio_ref_empate":  "Momio Referencia Empate (momio_ref_empate)",
    "momio_ref_visit":   "Momio Referencia Visitante (momio_ref_visit)",
    "p_imp_local":       "Probabilidad Implicita Local % (p_imp_local, ajust. vig)",
    "p_imp_empate":      "Probabilidad Implicita Empate % (p_imp_empate)",
    "p_imp_visit":       "Probabilidad Implicita Visitante % (p_imp_visit)",
    "n_bookmakers":      "Numero de Casas con Odds (n_bookmakers)",
    "casas_referencia":  "Casas de Apuestas Referencia (casas_referencia)",
    "ev_local":          "Expected Value Local (ev_local = P_modelo*momio - 1)",
    "ev_empate":         "Expected Value Empate (ev_empate)",
    "ev_visitante":      "Expected Value Visitante (ev_visitante)",
    "es_value_bet":      "Es Value Bet (ev > 0)",
    # ── Contexto enriquecido (Dixon-Coles + Motivación + Contexto) ──
    "motivacion_local":    "Motivacion Local (motivacion_local)",
    "motivacion_icon_h":   "Icono Motivacion Local (motivacion_icon_h)",
    "factor_motiv_h":      "Factor Motivacion Local (factor_motiv_h)",
    "motivacion_visita":   "Motivacion Visitante (motivacion_visita)",
    "motivacion_icon_a":   "Icono Motivacion Visitante (motivacion_icon_a)",
    "factor_motiv_a":      "Factor Motivacion Visitante (factor_motiv_a)",
    "alerta_empate":       "Alerta Empate Probable (alerta_empate)",
    "narrativa":           "Narrativa del Partido (narrativa)",
    "corners_h_est":       "Corners Estimados Local (corners_h_est)",
    "corners_a_est":       "Corners Estimados Visitante (corners_a_est)",
    "corners_total_rango": "Rango Corners Totales (corners_total_rango)",
    # FBref corners y tarjetas
    "corners_h_pred":      "Corners Predichos Local (corners_h)",
    "corners_a_pred":      "Corners Predichos Visitante (corners_a)",
    "corners_total_pred":  "Corners Total Predichos (corners_total)",
    "corners_fuente":      "Fuente Corners (FBref o lambda)",
    "amarillas_local":     "Amarillas Promedio Local (amarillas_local)",
    "amarillas_visita":    "Amarillas Promedio Visitante (amarillas_visita)",
    "amarillas_total":     "Amarillas Total Estimadas (amarillas_total)",
    "goleadores_local":    "Goleadores Local (goleadores_local)",
    "goleadores_visita":   "Goleadores Visitante (goleadores_visita)",
    "lambda_h_base":       "Lambda Local sin Motivacion (lambda_h_base)",
    "lambda_a_base":       "Lambda Visitante sin Motivacion (lambda_a_base)",
}


# ─── Generación del documento .docx ─────────────────────────────────────────────
def _narrative(row: dict) -> str:
    """Genera explicación en lenguaje natural de una predicción."""
    home    = row.get("local", "")
    away    = row.get("visitante", "")
    pred    = row.get("prediccion", "")
    p_loc   = row.get("p_local_pct", 0)
    p_emp   = row.get("p_empate_pct", 0)
    p_vis   = row.get("p_visit_pct", 0)
    lh      = row.get("lambda_h", 0)
    la      = row.get("lambda_a", 0)
    forma_h = row.get("forma_h", "N/D")
    forma_a = row.get("forma_a", "N/D")
    pts_h   = row.get("pts_forma_h", 0)
    pts_a   = row.get("pts_forma_a", 0)
    pos_h   = row.get("pos_tabla_h")
    pos_a   = row.get("pos_tabla_v")
    h2h_n   = row.get("h2h_n", 0)
    h2h_wh  = row.get("h2h_w_h", 0)
    h2h_d   = row.get("h2h_d", 0)
    h2h_wa  = row.get("h2h_w_a", 0)
    att_h   = row.get("att_h")
    def_a   = row.get("def_a")
    confianza = row.get("nivel_confianza", "BAJA")
    ventaja = row.get("ventaja", 0)

    pos_text_h = f"#{pos_h} en tabla" if pos_h else "posición desconocida"
    pos_text_a = f"#{pos_a} en tabla" if pos_a else "posición desconocida"

    att_txt = (f"ratio de ataque local de {att_h:.2f} ({'+' if att_h > 1 else ''}"
               f"{(att_h - 1)*100:.0f}% vs promedio liga)") if att_h else "datos insuficientes para rating"
    def_txt = (f"defensa visitante de {def_a:.2f} ({'+' if def_a > 1 else ''}"
               f"{(def_a - 1)*100:.0f}% vs promedio)") if def_a else "datos insuficientes"

    h2h_txt = (f"En los últimos {h2h_n} enfrentamientos directos: {home} ganó {h2h_wh}, "
               f"empates {h2h_d}, {away} ganó {h2h_wa}. ") if h2h_n > 0 else ""

    ev_loc = row.get("ev_local")
    ev_txt = (f" El Expected Value para victoria local es {ev_loc:+.3f} "
              f"({'VALUE BET' if ev_loc > 0 else 'sin valor'}).") if ev_loc is not None else ""

    if pred == "Local":
        text = (
            f"El modelo favorece a {home} como local (probabilidad {p_loc:.1f}%, "
            f"empate {p_emp:.1f}%, {away} {p_vis:.1f}%). "
            f"Se esperan {lh:.2f} goles de {home} y {la:.2f} de {away} "
            f"(ventaja {ventaja:+.2f} goles). "
            f"{home} ocupa el {pos_text_h} con forma reciente {forma_h} ({pts_h}/15 pts), "
            f"mientras {away} llega en {pos_text_a} con forma {forma_a} ({pts_a}/15 pts). "
            f"El {att_txt} y el ratio de {def_txt} justifican la preferencia por el local. "
            f"{h2h_txt}"
            f"Nivel de confianza del modelo: {confianza}.{ev_txt}"
        )
    elif pred == "Visitante":
        text = (
            f"A pesar de jugar fuera, {away} es favorito (probabilidad {p_vis:.1f}%, "
            f"empate {p_emp:.1f}%, {home} {p_loc:.1f}%). "
            f"El modelo espera {la:.2f} goles de {away} frente a {lh:.2f} de {home} "
            f"(diferencia {ventaja:+.2f} goles). "
            f"{away} ocupa el {pos_text_a} con forma {forma_a} ({pts_a}/15 pts), "
            f"contra {pos_text_h} de {home} con forma {forma_h} ({pts_h}/15 pts). "
            f"La calidad del visitante compensa la ventaja de local. "
            f"{h2h_txt}"
            f"Nivel de confianza del modelo: {confianza}.{ev_txt}"
        )
    else:  # Empate
        text = (
            f"El modelo detecta un equilibrio notable: {home} {p_loc:.1f}%, "
            f"empate {p_emp:.1f}%, {away} {p_vis:.1f}%. "
            f"Goles esperados prácticamente igualados: {home} {lh:.2f} — {away} {la:.2f} "
            f"(diferencia {ventaja:+.2f}). "
            f"Forma reciente similar: {home} {forma_h} ({pts_h} pts) vs {away} {forma_a} ({pts_a} pts). "
            f"{h2h_txt}"
            f"Nivel de confianza del modelo: {confianza}.{ev_txt}"
        )
    return text


def generate_docx(
    df: pd.DataFrame,
    ratings_info: dict,
    out_path: Path,
) -> None:
    """Genera el documento Word de metodología y predicciones."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  AVISO: python-docx no instalado. Ejecuta: pip install python-docx")
        return

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    VERDE    = RGBColor(0x2E, 0x86, 0x0E)
    AMARILLO = RGBColor(0xD9, 0x7A, 0x00)
    ROJO     = RGBColor(0xC0, 0x39, 0x2B)
    GRIS     = RGBColor(0x60, 0x60, 0x60)
    AZUL     = RGBColor(0x1A, 0x5F, 0x9B)
    CONF_COLORS = {"ALTA": VERDE, "MEDIA": AMARILLO, "BAJA": ROJO}

    mu_h     = ratings_info.get("mu_h", 1.40)
    mu_a     = ratings_info.get("mu_a", 1.18)
    n_hist   = ratings_info.get("n_hist", 0)
    n_equipos = ratings_info.get("n_equipos", 0)
    n_preds  = len(df)

    def _heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = AZUL
        return h


    # ── Portada ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_heading("DATA ANALYSIS PICKS", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.color.rgb = AZUL

    for txt, sz in [
        ("Predicciones y Metodología del Modelo", 14),
        ("Liga MX — Clausura 2026", 13),
        (f"Generado: {date.today().strftime('%d de %B de %Y')}", 11),
    ]:
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(sz)
        if sz == 13:
            p.runs[0].bold = True

    disc = doc.add_paragraph(
        "AVISO: Las predicciones son probabilísticas, no garantías de resultado. "
        "Las apuestas conllevan riesgo de pérdida."
    )
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc.runs[0].font.size = Pt(9)
    disc.runs[0].italic = True
    disc.runs[0].font.color.rgb = GRIS

    doc.add_page_break()

    # ── 1. Resumen Ejecutivo ────────────────────────────────────────────────────
    _heading("1. Resumen Ejecutivo")

    if n_preds > 0:
        dist     = df["prediccion"].value_counts().to_dict()
        alta_n   = (df["nivel_confianza"] == "ALTA").sum()
        media_n  = (df["nivel_confianza"] == "MEDIA").sum()
        baja_n   = (df["nivel_confianza"] == "BAJA").sum()
        value_n  = df["es_value_bet"].sum() if "es_value_bet" in df.columns else "N/A"
    else:
        dist = {}; alta_n = media_n = baja_n = 0; value_n = 0

    doc.add_paragraph(
        f"Este documento presenta {n_preds} predicciones para partidos de Liga MX "
        f"Clausura 2026, generadas mediante un modelo probabilístico de ratings "
        f"multiplicativos de ataque y defensa. El modelo se calibró con "
        f"{n_hist} partidos históricos de Liga MX de los últimos 150 días, "
        f"cubriendo {n_equipos} equipos.\n\n"
        f"Distribución de predicciones: "
        f"{dist.get('Local', 0)} victorias local · "
        f"{dist.get('Empate', 0)} empates · "
        f"{dist.get('Visitante', 0)} victorias visitante.\n\n"
        f"Nivel de confianza: {alta_n} ALTA · {media_n} MEDIA · {baja_n} BAJA.\n\n"
        f"Value Bets identificados: {value_n} (requiere integración de momios).\n\n"
        f"El modelo NO promete resultados garantizados. Su valor reside en "
        f"identificar ventajas estadísticas sostenidas a largo plazo sobre un "
        f"mínimo de 200 predicciones (benchmark de validación del MVP)."
    )

    doc.add_page_break()

    # ── 2. Metodología ─────────────────────────────────────────────────────────
    _heading("2. Metodología del Modelo")

    _heading("2.1 Fuente de Datos", 2)
    doc.add_paragraph(
        f"Los datos provienen de la API de Sofascore (sofascore6.p.rapidapi.com), "
        f"que proporciona resultados históricos con cobertura completa de Liga MX. "
        f"Se usan los últimos 150 días ({n_hist} partidos terminados) para calibrar "
        f"los ratings. La recolección es automática via collect_sofascore.py y se "
        f"actualiza en cada ejecución."
    )

    _heading("2.2 Ratings Multiplicativos Ataque/Defensa", 2)
    doc.add_paragraph(
        "El modelo calcula cuatro ratings por equipo, expresados como ratios "
        "respecto al promedio de la liga:"
    )
    items = [
        ("att_h (Ataque Local)",
         "Promedio ponderado de goles marcados en casa / μ_h. "
         "Valor > 1.0 = ataque local superior al promedio."),
        ("def_h (Defensa Local)",
         "Promedio ponderado de goles recibidos en casa / μ_a. "
         "Valor < 1.0 = defensa local superior al promedio."),
        ("att_a (Ataque Visitante)",
         "Promedio ponderado de goles marcados fuera / μ_a."),
        ("def_a (Defensa Visitante)",
         "Promedio ponderado de goles recibidos fuera / μ_h. "
         "Valor < 1.0 = defensa sólida fuera de casa."),
    ]
    for name, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    _heading("2.3 Cálculo de Goles Esperados (λ)", 2)
    doc.add_paragraph(
        "Los goles esperados por equipo se calculan combinando los ratings:"
    )
    for formula in [
        f"  λ_local   = att_h(local)  × def_a(visitante) × μ_h  [μ_h = {mu_h:.3f}]",
        f"  λ_visita  = att_a(visitante) × def_h(local)  × μ_a  [μ_a = {mu_a:.3f}]",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(formula)
        r.bold = True
        r.font.name = "Courier New"
        r.font.size = Pt(10)

    doc.add_paragraph(
        f"Todos los ratings se recortan al rango [0.30, 3.50] para evitar predicciones "
        f"extremas en equipos con pocos datos (mínimo {MIN_GAMES} partidos)."
    )

    _heading("2.4 Media Ponderada Reciente", 2)
    doc.add_paragraph(
        f"Los ratings se calculan con media ponderada de los últimos {WINDOW} partidos. "
        f"El partido más reciente tiene peso {WINDOW}, el más antiguo peso 1. "
        f"Esto captura la forma actual sin ignorar el historial completo."
    )

    _heading("2.5 Distribución de Poisson Independiente", 2)
    doc.add_paragraph(
        "Con λ_local y λ_visita se simula la distribución de probabilidad "
        "para todos los marcadores posibles (0 a 9 goles por equipo = 100 combinaciones):"
    )
    p = doc.add_paragraph()
    r = p.add_run("  P(k goles) = e^(−λ) × λ^k / k!")
    r.bold = True
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    doc.add_paragraph(
        "P(Local gana)    = Σ P(h,a) para todos los marcadores h > a\n"
        "P(Empate)        = Σ P(h,a) para todos los marcadores h = a\n"
        "P(Visitante gana)= Σ P(h,a) para todos los marcadores h < a\n\n"
        "Las tres probabilidades suman exactamente 1.0 por construcción del modelo."
    )

    _heading("2.6 Suavizado de Empate", 2)
    doc.add_paragraph(
        f"Si P(empate) < {DRAW_FLOOR*100:.0f}% tras el cálculo Poisson, se eleva a ese "
        f"umbral mínimo y el exceso se redistribuye proporcionalmente entre local y visitante. "
        f"El empate en fútbol mexicano tiene una frecuencia base real del ~25-30% que "
        f"la distribución de Poisson independiente tiende a subestimar."
    )

    _heading("2.7 Nivel de Confianza (Semáforo)", 2)
    doc.add_paragraph(
        "El semáforo refleja qué tan clara es la señal del modelo:"
    )
    for label, color, desc in [
        ("ALTA (verde)",   VERDE,    f"Diferencia entre mejor y 2a probabilidad > {ALTA_GAP*100:.0f}%. Señal clara y consistente."),
        ("MEDIA (naranja)", AMARILLO, f"Diferencia entre {MEDIA_GAP*100:.0f}% y {ALTA_GAP*100:.0f}%. Señal moderada."),
        ("BAJA (rojo)",    ROJO,     f"Diferencia < {MEDIA_GAP*100:.0f}%. Partido muy equilibrado, alta incertidumbre."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(label + ": ")
        r.bold = True
        r.font.color.rgb = color
        p.add_run(desc)

    doc.add_paragraph(
        f"Cuando se integren momios de casas de apuestas, el semáforo se basará en "
        f"Expected Value: ALTA si EV > {EV_ALTA*100:.0f}%, MEDIA si EV {EV_MEDIA*100:.0f}–{EV_ALTA*100:.0f}%."
    )

    _heading("2.8 Expected Value y Value Bets", 2)
    doc.add_paragraph(
        "El Expected Value cuantifica la ventaja estadística sobre el mercado:"
    )
    p = doc.add_paragraph()
    r = p.add_run("  EV = (P_modelo × momio_decimal) − 1")
    r.bold = True
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    doc.add_paragraph(
        "EV > 0  = value bet (el modelo detecta que el momio paga más de lo que debería)\n"
        "EV > 0.05 = value bet moderado  (pick MEDIA)\n"
        "EV > 0.10 = value bet alto      (pick ALTA)\n\n"
        "Ejemplo: Si el modelo asigna P(Local) = 55% y el momio es 2.10, "
        "entonces EV = 0.55 × 2.10 − 1 = +0.155 → value bet ALTA.\n\n"
        "IMPORTANTE: Un EV positivo es una ventaja estadística esperada a largo plazo, "
        "no una garantía de acierto en un partido individual."
    )

    _heading("2.9 Variables de Contexto", 2)
    doc.add_paragraph(
        "Además de los ratings base, el modelo incorpora variables de contexto que "
        "enriquecen la interpretación (no modifican directamente λ, pero permiten "
        "análisis cualitativo):"
    )
    ctx_items = [
        ("Posición en tabla",
         "Calculada desde todos los resultados disponibles de la temporada actual "
         "(puntos, diferencia de goles, goles a favor)."),
        ("H2H (Historial directo)",
         "Resultados de los últimos 5 enfrentamientos directos entre ambos equipos."),
        ("Forma reciente W/D/L",
         f"Últimos 5 partidos de cada equipo (independiente de local/visitante), "
         f"con puntos acumulados (W=3, D=1, L=0, máximo 15)."),
        ("Corners / Tiros de Esquina",
         "Promedio histórico de corners por partido desde FBref (passing_types.CK), "
         "diferenciado por rol: local (avg_corners_h) y visitante (avg_corners_a). "
         "Fallback: estimación via λ × 3.8 cuando FBref no tiene datos Liga MX."),
        ("Tarjetas Amarillas",
         "Promedio histórico de amarillas desde FBref (misc.CrdY), diferenciado "
         "por rol (local/visitante). Útil para mercados de tarjetas totales. "
         "Se muestra N/D si FBref no tiene datos para la temporada actual."),
    ]
    for name, desc in ctx_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ── 3. Glosario ────────────────────────────────────────────────────────────
    _heading("3. Glosario de Métricas del CSV")

    metrics = [
        ("Goles Esperados Local (lambda_h)",
         f"λ calculado como att_h × def_a_rival × μ_h ({mu_h:.2f}). Goles esperados del local."),
        ("Goles Esperados Visitante (lambda_a)",
         f"λ calculado como att_a × def_h_local × μ_a ({mu_a:.2f})."),
        ("Margen de Goles (ventaja)",
         "lambda_h − lambda_a. Positivo → favorito local. Negativo → favorito visitante."),
        ("att_h / att_a",
         "Ratio de ataque (local/visitante). >1.10 = ataque notablemente superior al promedio."),
        ("def_h / def_a",
         "Ratio de defensa. <0.90 = defensa notablemente superior al promedio."),
        ("Probabilidades %",
         "P(Local), P(Empate), P(Visitante). Siempre suman 100%. P(empate) >= 15% por suavizado."),
        ("Prob sin suavizar (raw)",
         "Probabilidades antes del suavizado de empate. Permite ver el impacto del ajuste."),
        ("Forma Reciente (W/D/L)",
         "Cadena de últimos 5 resultados. Ej: WWDLW = 2 vic, 1 emp, 1 der, 1 vic."),
        ("Puntos de Forma (pts_forma)",
         "Suma de puntos en últimos 5 partidos (máx. 15). >10 = buen momento, <5 = mal momento."),
        ("Posición en tabla",
         "Posición actual en Liga MX Clausura 2026 calculada desde resultados Sofascore."),
        ("H2H (h2h_w_h / h2h_d / h2h_w_a)",
         "Victorias local / Empates / Victorias visitante en los últimos 5 enfrentamientos directos."),
        ("Nivel de Confianza",
         "ALTA / MEDIA / BAJA. Con momios: basado en EV. Sin momios: basado en margen de probs."),
        ("Momio Referencia (momio_ref_*)",
         "Momio decimal de referencia por resultado (1=local, X=empate, 2=visitante)."),
        ("Expected Value (ev_*)",
         "EV = P_modelo × momio − 1. >0 = value bet. Requiere --odds-file para activar."),
        ("Es Value Bet",
         "TRUE si algún EV > 0 con los momios proporcionados."),
    ]

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light List Accent 1"
    for j, h in enumerate(["Métrica", "Descripción"]):
        tbl.rows[0].cells[j].text = h
        for para in tbl.rows[0].cells[j].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    for name, desc in metrics:
        row = tbl.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc
        for j in range(2):
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # ── 4. Predicciones por Partido ────────────────────────────────────────────
    _heading("4. Predicciones Detalladas")

    if df.empty:
        doc.add_paragraph(
            "No se encontraron partidos próximos de Liga MX en el horizonte consultado. "
            "Ejecuta collect_sofascore.py con --days 15 para actualizar los datos."
        )
    else:
        doc.add_paragraph(
            f"A continuación se presentan las {n_preds} predicciones con su razonamiento "
            f"estadístico. Las probabilidades suman 100% por diseño del modelo."
        )

        for i, (_, row) in enumerate(df.iterrows()):
            r = row.to_dict()
            home      = r.get("local", "")
            away      = r.get("visitante", "")
            fecha     = r.get("fecha", "")
            fase      = r.get("fase", "")
            pred      = r.get("prediccion", "")
            p_loc     = r.get("p_local_pct", 0)
            p_emp     = r.get("p_empate_pct", 0)
            p_vis     = r.get("p_visit_pct", 0)
            lh        = r.get("lambda_h", 0)
            la        = r.get("lambda_a", 0)
            confianza = r.get("nivel_confianza", "BAJA")

            # Título del partido
            h = doc.add_heading(f"{i+1}.  {home}  vs  {away}", 2)
            h.runs[0].font.color.rgb = AZUL

            # Fecha / fase
            p_info = doc.add_paragraph()
            r_info = p_info.add_run(f"Fecha: {fecha}   |   Jornada / Ronda: {fase or 'N/D'}")
            r_info.font.size = Pt(9)
            r_info.italic = True
            r_info.font.color.rgb = GRIS

            # Predicción + semáforo
            p_pred = doc.add_paragraph()
            r_pred = p_pred.add_run(f"Prediccion: {pred}   ")
            r_pred.bold = True
            r_pred.font.size = Pt(12)
            r_conf = p_pred.add_run(f"[Confianza: {confianza}]")
            r_conf.bold = True
            r_conf.font.size = Pt(12)
            r_conf.font.color.rgb = CONF_COLORS.get(confianza, GRIS)

            # Tabla de probabilidades
            prob_tbl = doc.add_table(rows=2, cols=3)
            prob_tbl.style = "Light Shading Accent 1"
            hdrs_p = [f"Local ({home[:15]})", "Empate", f"Visitante ({away[:15]})"]
            vals_p = [f"{p_loc:.1f}%", f"{p_emp:.1f}%", f"{p_vis:.1f}%"]
            for j in range(3):
                prob_tbl.rows[0].cells[j].text = hdrs_p[j]
                prob_tbl.rows[1].cells[j].text = vals_p[j]
                for para in prob_tbl.rows[0].cells[j].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                for para in prob_tbl.rows[1].cells[j].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(11)

            doc.add_paragraph()

            # Goles esperados
            p_g = doc.add_paragraph()
            p_g.add_run("Goles esperados: ").bold = True
            p_g.add_run(
                f"{home} {lh:.2f}  —  {away} {la:.2f}  "
                f"(ventaja {r.get('ventaja', 0):+.2f})"
            )

            # Forma
            p_f = doc.add_paragraph()
            p_f.add_run("Forma reciente: ").bold = True
            p_f.add_run(
                f"{home}: {r.get('forma_h', 'N/D')} ({r.get('pts_forma_h', 0)} pts)   |   "
                f"{away}: {r.get('forma_a', 'N/D')} ({r.get('pts_forma_a', 0)} pts)"
            )

            # Tabla de posiciones
            pos_h = r.get("pos_tabla_h")
            pos_a = r.get("pos_tabla_v")
            pts_th = r.get("pts_tabla_h")
            pts_ta = r.get("pts_tabla_v")
            if pos_h or pos_a:
                p_t = doc.add_paragraph()
                p_t.add_run("Posicion en tabla: ").bold = True
                p_t.add_run(
                    f"{home}: #{pos_h} ({pts_th} pts)   |   "
                    f"{away}: #{pos_a} ({pts_ta} pts)"
                )

            # H2H
            h2h_n = r.get("h2h_n", 0)
            if h2h_n > 0:
                p_h = doc.add_paragraph()
                p_h.add_run(f"H2H ultimos {h2h_n} partidos: ").bold = True
                p_h.add_run(
                    f"{home} {r.get('h2h_w_h', 0)}V / "
                    f"{r.get('h2h_d', 0)}E / "
                    f"{r.get('h2h_w_a', 0)}D "
                    f"(goles: {r.get('h2h_gf_h', 0)} - {r.get('h2h_gf_a', 0)})"
                )

            # Corners y Tarjetas (FBref)
            ck_h     = r.get("corners_h_pred")
            ck_a     = r.get("corners_a_pred")
            ck_tot   = r.get("corners_total_pred")
            ck_src   = r.get("corners_fuente", "N/D")
            amar_h   = r.get("amarillas_local")
            amar_a   = r.get("amarillas_visita")
            amar_tot = r.get("amarillas_total")
            if ck_h is not None:
                p_ck = doc.add_paragraph()
                p_ck.add_run("Corners estimados: ").bold = True
                p_ck.add_run(
                    f"{home[:16]}: ~{ck_h}  |  {away[:16]}: ~{ck_a}  "
                    f"(Total ~{ck_tot})  [Fuente: {ck_src}]"
                )
            if amar_h is not None or amar_a is not None:
                p_am = doc.add_paragraph()
                p_am.add_run("Amarillas (prom.): ").bold = True
                p_am.add_run(
                    f"{home[:16]}: {amar_h if amar_h is not None else 'N/D'}"
                    f"  |  {away[:16]}: {amar_a if amar_a is not None else 'N/D'}"
                    f"  (Total est.: {amar_tot if amar_tot is not None else 'N/D'})"
                )

            # EV
            ev_loc = r.get("ev_local")
            ev_emp = r.get("ev_empate")
            ev_vis = r.get("ev_visitante")
            if ev_loc is not None:
                p_ev = doc.add_paragraph()
                p_ev.add_run("Expected Value: ").bold = True
                p_ev.add_run(
                    f"Local {ev_loc:+.3f}  |  Empate {ev_emp:+.3f}  |  Visitante {ev_vis:+.3f}"
                )

            # Narrativa
            narr = _narrative(r)
            p_n = doc.add_paragraph()
            p_n.add_run("Analisis: ").bold = True
            r_n = p_n.add_run(narr)
            r_n.font.size = Pt(10)

            doc.add_paragraph()

    doc.add_page_break()

    # ── 5. Tabla Resumen ────────────────────────────────────────────────────────
    _heading("5. Tabla Resumen de Predicciones")

    if not df.empty:
        sum_cols = ["fecha", "local", "visitante", "prediccion",
                    "p_local_pct", "p_empate_pct", "p_visit_pct",
                    "lambda_h", "lambda_a", "nivel_confianza"]
        sum_hdrs = ["Fecha", "Local", "Visitante", "Prediccion",
                    "P.Local%", "P.Empate%", "P.Visit%", "λh", "λa", "Confianza"]

        tbl2 = doc.add_table(rows=1, cols=len(sum_cols))
        tbl2.style = "Light List Accent 1"
        for j, h in enumerate(sum_hdrs):
            tbl2.rows[0].cells[j].text = h
            for para in tbl2.rows[0].cells[j].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        for _, row in df.iterrows():
            r = row.to_dict()
            vals = [
                r.get("fecha", ""), r.get("local", "")[:18],
                r.get("visitante", "")[:18], r.get("prediccion", ""),
                f"{r.get('p_local_pct', 0):.1f}",
                f"{r.get('p_empate_pct', 0):.1f}",
                f"{r.get('p_visit_pct', 0):.1f}",
                f"{r.get('lambda_h', 0):.2f}",
                f"{r.get('lambda_a', 0):.2f}",
                r.get("nivel_confianza", ""),
            ]
            trow = tbl2.add_row()
            for j, v in enumerate(vals):
                trow.cells[j].text = str(v)
                for para in trow.cells[j].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

    doc.add_page_break()

    # ── 6. Aviso Legal ─────────────────────────────────────────────────────────
    _heading("6. Aviso Legal y Descargo de Responsabilidad")

    doc.add_paragraph(
        "Las predicciones contenidas en este documento son el resultado de modelos "
        "estadísticos automatizados. DATA ANALYSIS PICKS no garantiza que las "
        "predicciones sean correctas ni que su uso genere ganancias económicas.\n\n"
        "Las apuestas deportivas conllevan riesgo real de pérdida económica. "
        "Este documento es una herramienta de análisis estadístico con fines "
        "informativos exclusivamente.\n\n"
        "El producto NO se presenta como un servicio de apuestas garantizadas, "
        "seguras o sin riesgo. La ventaja estadística del modelo solo se manifiesta "
        "a largo plazo sobre un número significativo de predicciones "
        "(mínimo 200 picks para validación estadística según métricas del MVP).\n\n"
        "Conforme a la Ley Federal de Protección de Datos Personales en Posesión de "
        "los Particulares (LFPDPPP) de México, los datos de usuarios son tratados "
        "con confidencialidad y no son compartidos con terceros sin consentimiento.\n\n"
        "Las probabilidades implícitas en los momios de las casas de apuestas incluyen "
        "un margen operativo (vig/juice) que reduce el retorno esperado del apostador. "
        "El modelo solo recomienda picks donde P_modelo > P_implícita (EV positivo), "
        "compensando estadísticamente dicho margen."
    )

    doc.save(out_path)
    print(f"  Documento guardado: {out_path}")


# ─── Main ─────────────────────────────────────────────────────────────────────────
def main(odds_file: str | None = None, odds_key: str | None = None) -> None:
    print("=" * 70)
    print("  DATA ANALYSIS PICKS — Liga MX Clausura 2026")
    print("  Modelo: Ratings Multiplicativos + Poisson + Semaforo + EV")
    print("=" * 70)

    # ── 1. Cargar datos ──────────────────────────────────────────────────────────
    if not EVENTS_PATH.exists():
        print(f"\n  ERROR: {EVENTS_PATH} no encontrado.")
        print("  Ejecuta primero: python collect_sofascore.py --history-days 150")
        return

    print(f"\n[1/4] Cargando {EVENTS_PATH}...")
    df_all = pd.read_parquet(EVENTS_PATH)
    df_all["date"] = pd.to_datetime(df_all["date"])

    liga_df = df_all[df_all["tournament_id"] == LIGA_MX_ID].copy()
    print(f"  Liga MX historico: {len(liga_df)} partidos")

    hist_df = liga_df[liga_df["is_result"] == True].dropna(
        subset=["home_goals", "away_goals"]
    ).copy()

    # Partidos próximos: leer de sofascore_upcoming.parquet (guardado por collect_sofascore.py)
    upcoming = []
    if UPCOMING_PATH.exists():
        up_df = pd.read_parquet(UPCOMING_PATH)
        liga_up = up_df[up_df["tournament_id"] == LIGA_MX_ID]
        upcoming = liga_up.to_dict("records")
        print(f"  Proximos (upcoming.parquet): {len(upcoming)} partidos Liga MX")
    else:
        print("  AVISO: sofascore_upcoming.parquet no encontrado.")

    print(f"  Historico: {len(hist_df)} | Proximos: {len(upcoming)}")

    if hist_df.empty:
        print("  ERROR: No hay datos historicos. Re-ejecuta collect_sofascore.py.")
        return

    # ── 2. Calcular ratings, tabla, H2H ──────────────────────────────────────────
    print("\n[2/4] Calculando ratings y tabla...")
    ratings = compute_ratings(hist_df)
    table   = compute_table(hist_df)

    n_equipos = len(ratings["teams"])
    print(f"  Ratings: {n_equipos} equipos")
    print(f"  mu_h={ratings['mu_h']:.3f} | mu_a={ratings['mu_a']:.3f} | "
          f"home_adv={ratings['home_adv']:.3f}")
    if not table.empty:
        print(f"  Tabla calculada: {len(table)} equipos")

    # ── 3. Cargar odds ────────────────────────────────────────────────────────────
    # Prioridad: --odds-key (The Odds API) > --odds-file (JSON manual) > sin momios
    odds = None
    if odds_key:
        print(f"\n[3/4b] Obteniendo momios desde The Odds API...")
        odds = fetch_odds_ligamx(odds_key)
        if not odds:
            print("  Sin momios disponibles — predicciones sin EV.")
    elif odds_file:
        try:
            raw = json.loads(Path(odds_file).read_bytes().decode("utf-8"))
            # Convertir formato manual a formato interno (añadir home_api/away_api)
            odds = {}
            for key, v in raw.items():
                parts = key.split("_vs_")
                home_k = parts[0] if len(parts) == 2 else key
                away_k = parts[1] if len(parts) == 2 else ""
                v.setdefault("home_api", home_k)
                v.setdefault("away_api", away_k)
                odds[key] = v
            print(f"\n  Momios cargados: {len(odds)} partidos desde {odds_file}")
        except Exception as e:
            print(f"\n  AVISO: No se pudo cargar odds file: {e}")

    # ── 3.5. Cargar datos FBref (corners + tarjetas) ─────────────────────────────
    print(f"\n[2.5/4] Estadísticas FBref Liga MX (corners + tarjetas amarillas)...")
    fbref_stats = load_fbref_ligamx_stats()
    n_fbref = len(fbref_stats)
    if n_fbref:
        print(f"  FBref: {n_fbref} equipos con datos históricos")
    else:
        print("  FBref: sin datos — corners se estimarán por λ, tarjetas N/D")

    # ── 4. Generar predicciones ───────────────────────────────────────────────────
    print(f"\n[3/4] Generando predicciones para {len(upcoming)} partidos proximos...")

    if not upcoming:
        print("  AVISO: No hay partidos proximos en el parquet.")
        print("  Ejecuta: python collect_sofascore.py --days 15")
        # Generamos docx vacío de metodología de todas formas
        predictions_df = pd.DataFrame()
    else:
        predictions_df = generate_predictions(
            upcoming, hist_df, ratings, table, odds, fbref_stats=fbref_stats
        )

    # ── 5. Guardar CSV ────────────────────────────────────────────────────────────
    print(f"\n[4/4] Guardando salidas...")

    if not predictions_df.empty:
        out_df = predictions_df.rename(columns=COL_NAMES)
        out_df.to_csv(OUT_CSV, index=False, encoding="utf-8-sig")
        print(f"  CSV: {OUT_CSV}  ({len(out_df)} partidos x {len(out_df.columns)} columnas)")

        dist = predictions_df["prediccion"].value_counts().to_dict()
        alta = (predictions_df["nivel_confianza"] == "ALTA").sum()
        media = (predictions_df["nivel_confianza"] == "MEDIA").sum()
        baja = (predictions_df["nivel_confianza"] == "BAJA").sum()
        print(f"  Distribucion: {dist}")
        print(f"  Confianza: {alta} ALTA | {media} MEDIA | {baja} BAJA")

        SEP = "-" * 70
        print()
        print(SEP)
        for _, row in predictions_df.iterrows():
            r = row.to_dict()
            ev_str = ""
            evs = {k: r[k] for k in ("ev_local", "ev_empate", "ev_visitante") if r.get(k) is not None}
            if evs:
                best_ev_label = max(evs, key=evs.get)
                best_ev_val   = evs[best_ev_label]
                ev_str = f"  EV={best_ev_val:+.3f}({'VB' if best_ev_val > 0 else ''})"
            print(
                f"  {r['fecha']}  {r['local'][:18]:<18} vs  {r['visitante'][:18]:<18}"
                f"  [{r['prediccion']:<10}]  {r['p_local_pct']:>5.1f}%/"
                f"{r['p_empate_pct']:>4.1f}%/{r['p_visit_pct']:>4.1f}%"
                f"  [{r['nivel_confianza']}]{ev_str}"
            )
        print(SEP)

    # ── 6. Generar DOCX ───────────────────────────────────────────────────────────
    ratings_info = {
        "mu_h":      ratings["mu_h"],
        "mu_a":      ratings["mu_a"],
        "home_adv":  ratings["home_adv"],
        "n_hist":    len(hist_df),
        "n_equipos": n_equipos,
    }
    generate_docx(predictions_df, ratings_info, OUT_DOCX)
    print()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Liga MX — Predicciones + Value Bets + Metodologia (Data Analysis Picks)"
    )
    parser.add_argument(
        "--odds-key",
        type=str,
        default=None,
        metavar="API_KEY",
        help=(
            "Clave de The Odds API para obtener momios en tiempo real. "
            "Registro gratuito (500 req/mes) en https://the-odds-api.com"
        ),
    )
    parser.add_argument(
        "--odds-file",
        type=str,
        default=None,
        metavar="FILE",
        help=(
            "JSON con momios manuales: "
            "{\"Club Puebla_vs_Club Necaxa\": {\"local\": 2.80, \"empate\": 3.10, \"visitante\": 2.40}}"
        ),
    )
    args = parser.parse_args()
    main(odds_file=args.odds_file, odds_key=args.odds_key)
